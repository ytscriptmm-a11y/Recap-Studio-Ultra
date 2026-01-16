import streamlit as st
import google.generativeai as genai
import time
import os
import tempfile
import gc
import io
import hashlib
import asyncio
import struct
import re
import json
from datetime import datetime, date
from PIL import Image

# --- LIBRARY IMPORTS WITH FALLBACKS ---
PDF_AVAILABLE = True
DOCX_AVAILABLE = True
GDOWN_AVAILABLE = True
SUPABASE_AVAILABLE = True
EDGE_TTS_AVAILABLE = True
GENAI_NEW_AVAILABLE = True
YT_DLP_AVAILABLE = True

try:
    import PyPDF2
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
except ImportError:
    DOCX_AVAILABLE = False

try:
    import gdown
except ImportError:
    GDOWN_AVAILABLE = False

try:
    from supabase import create_client
except ImportError:
    SUPABASE_AVAILABLE = False

try:
    import edge_tts
except ImportError:
    EDGE_TTS_AVAILABLE = False

try:
    from google import genai as genai_new
    from google.genai import types
except ImportError:
    GENAI_NEW_AVAILABLE = False

try:
    import yt_dlp
except ImportError:
    YT_DLP_AVAILABLE = False

# Supabase setup
SUPABASE_URL = "https://ohjvgupjocgsirhwuobf.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Im9oanZndXBqb2Nnc2lyaHd1b2JmIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NjU5MzkwMTgsImV4cCI6MjA4MTUxNTAxOH0.oZxQZ6oksjbmEeA_m8c44dG_z5hHLwtgoJssgK2aogI"
supabase = None
if SUPABASE_AVAILABLE:
    try:
        supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    except:
        SUPABASE_AVAILABLE = False

# Daily limits
DAILY_LIMITS = {
    'content': 10,
    'translate': 10,
    'tts_gemini': 10,
    'thumbnail': 10
}

st.set_page_config(
    page_title="AI Studio Pro",
    layout="centered",
    initial_sidebar_state="collapsed",
    page_icon="🎬"
)

# === MODERN UI CSS ===
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Myanmar:wght@300;400;500;600;700&family=Poppins:wght@300;400;500;600;700&display=swap');

:root {
    --primary: #6366f1;
    --primary-light: #818cf8;
    --primary-dark: #4f46e5;
    --accent: #22d3ee;
    --accent-pink: #f472b6;
    --bg-dark: #0a0a1a;
    --bg-glass: rgba(255, 255, 255, 0.05);
    --text-primary: #f1f5f9;
    --text-secondary: #94a3b8;
    --border-glass: rgba(255, 255, 255, 0.1);
    --success: #10b981;
    --warning: #f59e0b;
    --error: #ef4444;
}

* { font-family: 'Poppins', 'Noto Sans Myanmar', sans-serif !important; }

.stApp {
    background: linear-gradient(135deg, var(--bg-dark) 0%, #0f172a 50%, #1e1b4b 100%) !important;
}

.stApp::before {
    content: '';
    position: fixed;
    top: 0; left: 0; right: 0; bottom: 0;
    background: 
        radial-gradient(circle at 20% 80%, rgba(99, 102, 241, 0.15) 0%, transparent 50%),
        radial-gradient(circle at 80% 20%, rgba(34, 211, 238, 0.1) 0%, transparent 50%);
    pointer-events: none;
    z-index: 0;
}

header, #MainMenu, footer, [data-testid="stDecoration"] { display: none !important; }

[data-testid="block-container"] {
    max-width: 100% !important;
    padding: 1rem !important;
}

@media (min-width: 768px) {
    [data-testid="block-container"] { max-width: 900px !important; padding: 2rem !important; }
}

div[data-testid="stVerticalBlockBorderWrapper"] {
    background: var(--bg-glass) !important;
    backdrop-filter: blur(20px) !important;
    border: 1px solid var(--border-glass) !important;
    border-radius: 16px !important;
    padding: 1.2rem !important;
    margin-bottom: 0.8rem !important;
}

.stTextInput > div > div > input,
.stTextArea > div > div > textarea {
    background: rgba(15, 23, 42, 0.6) !important;
    color: var(--text-primary) !important;
    border: 1px solid var(--border-glass) !important;
    border-radius: 10px !important;
    padding: 10px 14px !important;
}

.stButton > button {
    background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%) !important;
    color: white !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 10px 20px !important;
    font-weight: 600 !important;
    transition: all 0.3s ease !important;
}

.stButton > button:hover { transform: translateY(-2px) !important; }

.stDownloadButton > button {
    background: linear-gradient(135deg, var(--success) 0%, #059669 100%) !important;
}

.stTabs [data-baseweb="tab-list"] {
    background: var(--bg-glass) !important;
    padding: 6px !important;
    border-radius: 12px !important;
    gap: 4px !important;
    flex-wrap: wrap !important;
    justify-content: center !important;
}

.stTabs [data-baseweb="tab"] {
    color: var(--text-secondary) !important;
    border-radius: 8px !important;
    padding: 8px 12px !important;
    font-size: 12px !important;
}

.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, var(--primary) 0%, var(--accent) 100%) !important;
    color: white !important;
}

h1, h2, h3 { color: var(--text-primary) !important; }
h1 {
    background: linear-gradient(135deg, var(--primary-light), var(--accent)) !important;
    -webkit-background-clip: text !important;
    -webkit-text-fill-color: transparent !important;
}

p, span, label { color: var(--text-secondary) !important; }

.stSelectbox > div > div {
    background: rgba(15, 23, 42, 0.6) !important;
    border: 1px solid var(--border-glass) !important;
    border-radius: 10px !important;
}

div[data-testid="stFileUploader"] section {
    background: rgba(15, 23, 42, 0.4) !important;
    border: 2px dashed var(--border-glass) !important;
    border-radius: 12px !important;
}

[data-testid="stMetricValue"] { color: var(--accent) !important; }

hr { background: var(--border-glass) !important; height: 1px !important; border: none !important; }

.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--primary), var(--accent)) !important;
}

.stRadio > div > label {
    background: var(--bg-glass) !important;
    padding: 8px 12px !important;
    border-radius: 8px !important;
    border: 1px solid var(--border-glass) !important;
}

audio { width: 100% !important; border-radius: 10px !important; }

.usage-box {
    background: rgba(99, 102, 241, 0.1);
    border: 1px solid rgba(99, 102, 241, 0.3);
    border-radius: 10px;
    padding: 10px 15px;
    margin: 10px 0;
}

.limit-warning {
    background: rgba(239, 68, 68, 0.1);
    border: 1px solid rgba(239, 68, 68, 0.3);
    border-radius: 10px;
    padding: 10px 15px;
    color: #fca5a5;
}
</style>
""", unsafe_allow_html=True)

# === HELPER FUNCTIONS ===
def get_app_api_key():
    """Get APP API key from Streamlit secrets"""
    try:
        return st.secrets["google"]["app_api_key"]
    except:
        return None

def parse_mime(m):
    b, r = 16, 24000
    for p in m.split(";"):
        p = p.strip()
        if p.lower().startswith("rate="):
            r = int(p.split("=")[1])
        elif p.startswith("audio/L"):
            b = int(p.split("L")[1])
    return b, r

def to_wav(d, m):
    b, r = parse_mime(m)
    h = struct.pack("<4sI4s4sIHHIIHH4sI", b"RIFF", 36+len(d), b"WAVE", b"fmt ", 16, 1, 1, r, r*b//8, b//8, b, b"data", len(d))
    return h + d

def cleanup():
    gc.collect()

def get_text(r):
    try:
        if not r or not r.candidates:
            return None, "No response"
        parts = r.candidates[0].content.parts if hasattr(r.candidates[0], 'content') else []
        t = "\n".join([p.text for p in parts if hasattr(p, 'text') and p.text])
        return (t, None) if t else (None, "No text")
    except Exception as e:
        return None, str(e)

def call_api(m, c, to=900):
    for i in range(3):
        try:
            r = m.generate_content(c, request_options={"timeout": to})
            t, e = get_text(r)
            if t:
                return r, None
            if i < 2:
                time.sleep(10)
        except Exception as e:
            err_str = str(e).lower()
            if any(x in err_str for x in ['rate', 'quota', '429']):
                if i < 2:
                    time.sleep(10 * (2**i))
                else:
                    return None, "Rate limit ပြည့်သွားပါပြီ။ ခဏစောင့်ပြီး ပြန်ကြိုးစားပါ။"
            else:
                return None, str(e)
    return None, "အကြိမ်ရေ ပြည့်သွားပါပြီ။ နောက်မှ ပြန်ကြိုးစားပါ။"

def upload_gem(p, s=None):
    try:
        if s:
            s.info(f"📤 Uploading ({os.path.getsize(p)/(1024*1024):.1f}MB)...")
        f = genai.upload_file(p)
        w = 0
        while f.state.name == "PROCESSING":
            w += 1
            if s:
                s.info(f"⏳ Processing...({w*2}s)")
            time.sleep(2)
            f = genai.get_file(f.name)
            if w > 300:
                return None
        return f if f.state.name != "FAILED" else None
    except Exception as e:
        if s:
            s.error(str(e))
        return None

def save_up(u):
    try:
        ext = u.name.split('.')[-1] if '.' in u.name else 'mp4'
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}")
        u.seek(0, 2)
        sz = u.tell()
        u.seek(0)
        prog = st.progress(0)
        wr = 0
        while ch := u.read(10*1024*1024):
            tmp.write(ch)
            wr += len(ch)
            prog.progress(min(wr/sz, 1.0))
        tmp.close()
        prog.empty()
        return tmp.name, None
    except Exception as e:
        return None, str(e)

def rm_file(p):
    if p and os.path.exists(p):
        try:
            os.remove(p)
        except:
            pass

def read_file(u):
    try:
        t = u.type
        if t == "text/plain":
            return u.getvalue().decode("utf-8")
        elif t == "application/pdf" and PDF_AVAILABLE:
            return "\n".join([p.extract_text() or "" for p in PyPDF2.PdfReader(io.BytesIO(u.getvalue())).pages])
        elif "wordprocessingml" in t and DOCX_AVAILABLE:
            return "\n".join([p.text for p in Document(io.BytesIO(u.getvalue())).paragraphs])
        return None
    except:
        return None

def get_gid(url):
    try:
        if 'drive.google.com' in url:
            if '/file/d/' in url:
                return url.split('/file/d/')[1].split('/')[0].split('?')[0]
            elif 'id=' in url:
                return url.split('id=')[1].split('&')[0]
        return None
    except:
        return None

def dl_gdrive(url, s=None):
    if not GDOWN_AVAILABLE:
        return None, "gdown မရနိုင်ပါ"
    try:
        fid = get_gid(url)
        if not fid:
            return None, "Google Drive link မှားနေပါတယ်"
        if s:
            s.info("📥 Google Drive မှ ဒေါင်းလုဒ်လုပ်နေပါတယ်...")
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4").name
        if gdown.download(f"https://drive.google.com/uc?id={fid}", tmp, quiet=True, fuzzy=True):
            if os.path.exists(tmp) and os.path.getsize(tmp) > 1000:
                return tmp, None
        return None, "ဒေါင်းလုဒ် မအောင်မြင်ပါ"
    except Exception as e:
        return None, str(e)

def download_video_url(url, status=None):
    """Download from YouTube, TikTok, Facebook - NO COOKIES"""
    if 'drive.google.com' in url:
        return dl_gdrive(url, status)
    
    if not YT_DLP_AVAILABLE:
        return None, "yt-dlp မရနိုင်ပါ။ Google Drive link သို့မဟုတ် File upload သုံးပါ။"
    
    try:
        if status:
            status.info("📥 Video ဒေါင်းလုဒ်လုပ်နေပါတယ်...")
        
        output_path = f"/tmp/video_{int(time.time())}.mp4"
        ydl_opts = {
            'format': 'best[ext=mp4]/best',
            'outtmpl': output_path,
            'quiet': True,
            'no_warnings': True,
            'socket_timeout': 60,
            'noplaylist': True,
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        
        if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
            return output_path, None
        else:
            return None, "ဒေါင်းလုဒ် မအောင်မြင်ပါ"
    except Exception as e:
        err_msg = str(e).lower()
        if 'sign in' in err_msg or 'age' in err_msg:
            return None, """⚠️ ဒီ video ကို ဒေါင်းလုဒ်လုပ်လို့ မရပါ။

အကြောင်းရင်း: Age-restricted သို့မဟုတ် Login လိုအပ်သော video ဖြစ်နိုင်ပါတယ်။

💡 အခြားနည်းလမ်းများ:
1. Video ကို ကိုယ်တိုင် ဒေါင်းလုဒ်လုပ်ပြီး File Upload သုံးပါ
2. Google Drive မှာ တင်ပြီး link ပေးပါ"""
        return None, f"ဒေါင်းလုဒ် မအောင်မြင်ပါ: {str(e)[:100]}"

def hash_pw(p):
    return hashlib.sha256(p.encode()).hexdigest()

def login(e, p):
    if not supabase:
        return None, "Database ချိတ်ဆက်မှု မအောင်မြင်ပါ"
    try:
        r = supabase.table('users').select('*').eq('email', e).eq('password', hash_pw(p)).execute()
        if r.data:
            u = r.data[0]
            return (u, "OK") if u['approved'] else (None, "Admin approval စောင့်ဆိုင်းနေပါသည်")
        return None, "Email သို့မဟုတ် Password မှားနေပါတယ်"
    except Exception as ex:
        return None, str(ex)

def register(e, p):
    if not supabase:
        return False, "Database ချိတ်ဆက်မှု မအောင်မြင်ပါ"
    try:
        if supabase.table('users').select('email').eq('email', e).execute().data:
            return False, "ဒီ Email ကို အသုံးပြုပြီးသားဖြစ်နေပါတယ်"
        supabase.table('users').insert({
            "email": e,
            "password": hash_pw(p),
            "approved": False,
            "is_admin": False,
            "usage_content": 0,
            "usage_translate": 0,
            "usage_tts": 0,
            "usage_thumbnail": 0,
            "last_usage_date": str(date.today())
        }).execute()
        return True, "အကောင့်ဖွင့်ပြီးပါပြီ! Admin approval စောင့်ပါ။"
    except Exception as ex:
        return False, str(ex)

def get_usage(user_id):
    """Get user's daily usage, reset if new day"""
    if not supabase:
        return {'content': 0, 'translate': 0, 'tts': 0, 'thumbnail': 0}
    
    try:
        r = supabase.table('users').select('*').eq('id', user_id).execute()
        if r.data:
            u = r.data[0]
            last_date = u.get('last_usage_date', '')
            today = str(date.today())
            
            # Reset if new day
            if last_date != today:
                supabase.table('users').update({
                    'usage_content': 0,
                    'usage_translate': 0,
                    'usage_tts': 0,
                    'usage_thumbnail': 0,
                    'last_usage_date': today
                }).eq('id', user_id).execute()
                return {'content': 0, 'translate': 0, 'tts': 0, 'thumbnail': 0}
            
            return {
                'content': u.get('usage_content', 0) or 0,
                'translate': u.get('usage_translate', 0) or 0,
                'tts': u.get('usage_tts', 0) or 0,
                'thumbnail': u.get('usage_thumbnail', 0) or 0
            }
    except:
        pass
    return {'content': 0, 'translate': 0, 'tts': 0, 'thumbnail': 0}

def increment_usage(user_id, feature):
    """Increment usage count for a feature"""
    if not supabase:
        return
    try:
        col = f'usage_{feature}'
        r = supabase.table('users').select(col).eq('id', user_id).execute()
        if r.data:
            current = r.data[0].get(col, 0) or 0
            supabase.table('users').update({
                col: current + 1,
                'last_usage_date': str(date.today())
            }).eq('id', user_id).execute()
    except:
        pass

def check_limit(user_id, feature, api_type):
    """Check if user has reached daily limit"""
    if api_type == 'own':
        return True, 0  # Unlimited for OWN API
    
    usage = get_usage(user_id)
    current = usage.get(feature, 0)
    limit = DAILY_LIMITS.get(feature, 10)
    
    if current >= limit:
        return False, 0
    return True, limit - current

def srt_to_text(srt_content):
    lines = srt_content.split('\n')
    text_lines = []
    for line in lines:
        line = line.strip()
        if not line or line.isdigit() or '-->' in line:
            continue
        text_lines.append(line)
    return '\n'.join(text_lines)

def text_to_srt(text, sec_per_line=3):
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    srt_out = []
    for i, line in enumerate(lines):
        start = i * sec_per_line
        end = (i + 1) * sec_per_line
        sh, sm, ss = start // 3600, (start % 3600) // 60, start % 60
        eh, em, es = end // 3600, (end % 3600) // 60, end % 60
        srt_out.append(f"{i+1}")
        srt_out.append(f"{sh:02d}:{sm:02d}:{ss:02d},000 --> {eh:02d}:{em:02d}:{es:02d},000")
        srt_out.append(line)
        srt_out.append("")
    return '\n'.join(srt_out)

# === TTS FUNCTIONS ===
def edge_voices():
    return {
        "🇲🇲 Myanmar - Thiha (ကျား)": "my-MM-ThihaNeural",
        "🇲🇲 Myanmar - Nilar (မ)": "my-MM-NilarNeural",
        "🇺🇸 English - Jenny (မ)": "en-US-JennyNeural",
        "🇺🇸 English - Guy (ကျား)": "en-US-GuyNeural",
        "🇺🇸 English - Aria (မ)": "en-US-AriaNeural",
        "🇺🇸 English - Davis (ကျား)": "en-US-DavisNeural",
        "🇬🇧 British - Sonia (မ)": "en-GB-SoniaNeural",
        "🇬🇧 British - Ryan (ကျား)": "en-GB-RyanNeural",
        "🇹🇭 Thai - Premwadee (မ)": "th-TH-PremwadeeNeural",
        "🇹🇭 Thai - Niwat (ကျား)": "th-TH-NiwatNeural",
        "🇨🇳 Chinese - Xiaoxiao (မ)": "zh-CN-XiaoxiaoNeural",
        "🇨🇳 Chinese - Yunyang (ကျား)": "zh-CN-YunyangNeural",
        "🇯🇵 Japanese - Nanami (မ)": "ja-JP-NanamiNeural",
        "🇯🇵 Japanese - Keita (ကျား)": "ja-JP-KeitaNeural",
        "🇰🇷 Korean - SunHi (မ)": "ko-KR-SunHiNeural",
        "🇰🇷 Korean - InJoon (ကျား)": "ko-KR-InJoonNeural",
        "🇮🇳 Hindi - Swara (မ)": "hi-IN-SwaraNeural",
        "🇮🇳 Hindi - Madhur (ကျား)": "hi-IN-MadhurNeural",
        "🇻🇳 Vietnamese - HoaiMy (မ)": "vi-VN-HoaiMyNeural",
        "🇻🇳 Vietnamese - NamMinh (ကျား)": "vi-VN-NamMinhNeural",
    }

def gemini_voices():
    return {
        "Puck (ကျား)": "Puck",
        "Charon (ကျား)": "Charon",
        "Kore (မ)": "Kore",
        "Fenrir (ကျား)": "Fenrir",
        "Aoede (မ)": "Aoede",
        "Leda (မ)": "Leda",
        "Orus (ကျား)": "Orus",
        "Zephyr (ကျား)": "Zephyr",
        "Helios (ကျား)": "Helios",
        "Perseus (ကျား)": "Perseus",
        "Callirrhoe (မ)": "Callirrhoe",
        "Autonoe (မ)": "Autonoe",
        "Enceladus (ကျား)": "Enceladus",
        "Iapetus (ကျား)": "Iapetus",
        "Umbriel (ကျား)": "Umbriel",
        "Algieba (မ)": "Algieba",
        "Despina (မ)": "Despina",
        "Erinome (မ)": "Erinome",
        "Gacrux (ကျား)": "Gacrux",
        "Achird (ကျား)": "Achird",
        "Zubenelgenubi (ကျား)": "Zubenelgenubi",
        "Schedar (မ)": "Schedar",
        "Sadachbia (ကျား)": "Sadachbia",
        "Sadaltager (ကျား)": "Sadaltager",
        "Sulafat (မ)": "Sulafat"
    }

def voice_styles():
    return {
        "🎬 ပုံမှန် ဇာတ်လမ်းပြန်ပြောခြင်း": "Narrate in an engaging storytelling style.",
        "🔥 သည်းထိတ်ရင်ဖို": "Dramatic and suspenseful narration, serious and intense.",
        "😊 ပေါ့ပေါ့ပါးပါး": "Casual, friendly, energetic manner like a YouTuber.",
        "🎃 ထိတ်လန့်စရာ": "Chilling, eerie tone for horror content.",
        "🎭 ခံစားချက်ပြည့်": "Deep emotional expression, dramatic reading.",
        "📺 သတင်းကြေငြာ": "Professional news anchor style.",
        "🎓 မှတ်တမ်းရုပ်ရှင်": "Calm, educational documentary style.",
        "🎪 စိတ်ကြိုက်": ""
    }

def gen_edge(txt, v, r=0):
    if not EDGE_TTS_AVAILABLE:
        return None, "Edge TTS မရနိုင်ပါ"
    try:
        out = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3").name
        rs = f"+{r}%" if r >= 0 else f"{r}%"
        async def _g():
            await edge_tts.Communicate(txt, v, rate=rs).save(out)
        asyncio.run(_g())
        return out, None
    except Exception as e:
        return None, str(e)

def gen_gemini_tts(key, txt, v, mdl, style="", speed=1.0):
    if not GENAI_NEW_AVAILABLE:
        return None, "Gemini TTS မရနိုင်ပါ"
    try:
        cl = genai_new.Client(api_key=key)
        speed_inst = ""
        if speed < 1.0:
            speed_inst = f" Speak slowly at {speed}x speed."
        elif speed > 1.0:
            speed_inst = f" Speak faster at {speed}x speed."
        
        full_text = f"[Style: {style}{speed_inst}]\n\n{txt}" if style else txt
        
        cfg = types.GenerateContentConfig(
            temperature=1,
            response_modalities=["audio"],
            speech_config=types.SpeechConfig(
                voice_config=types.VoiceConfig(
                    prebuilt_voice_config=types.PrebuiltVoiceConfig(voice_name=v)
                )
            )
        )
        
        aud = b""
        mime = "audio/L16;rate=24000"
        for ch in cl.models.generate_content_stream(
            model=mdl,
            contents=[types.Content(role="user", parts=[types.Part.from_text(text=full_text)])],
            config=cfg
        ):
            if ch.candidates and ch.candidates[0].content and ch.candidates[0].content.parts:
                p = ch.candidates[0].content.parts[0]
                if hasattr(p, 'inline_data') and p.inline_data and p.inline_data.data:
                    aud += p.inline_data.data
                    mime = p.inline_data.mime_type
        
        if not aud:
            return None, "Audio မထွက်ပါ"
        
        out = tempfile.NamedTemporaryFile(delete=False, suffix=".wav").name
        with open(out, "wb") as f:
            f.write(to_wav(aud, mime))
        return out, None
    except Exception as e:
        return None, str(e)

# === CONTENT GENERATION ===
def get_content_types():
    return {
        "📰 ဆောင်းပါး": "article",
        "🏆 အောင်မြင်ရေး": "success",
        "📖 ဝတ္ထုတို": "story",
        "🧒 ပုံပြင်": "tale",
        "📢 သတင်း": "news",
        "🎬 ဇာတ်လမ်း": "drama",
        "👻 သရဲဇာတ်လမ်း": "horror",
        "💔 ဂမ္ဘီရ": "tragic",
        "💕 အချစ်ဇာတ်လမ်း": "romance",
        "🔮 စိတ်ကူးယဉ်": "fantasy",
        "🔍 လျှို့ဝှက်ဆန်းကြယ်": "mystery",
        "😂 ဟာသ": "comedy",
        "💪 လှုံ့ဆော်စာ": "motivational",
        "📚 ပညာရေး": "educational",
        "🎯 စိတ်ကြိုက်": "custom"
    }

def get_tones():
    return {
        "📝 ပုံမှန်": "",
        "😊 ပေါ့ပေါ့ပါးပါး": "ပေါ့ပေါ့ပါးပါး၊ ဖတ်ရလွယ်ကူစွာ ရေးပါ။",
        "🎭 ဂမ္ဘီရ": "ဂမ္ဘီရဆန်ဆန်၊ လေးနက်စွာ ရေးပါ။",
        "🔥 Gen Z": "Gen Z လူငယ်တွေ စိတ်ဝင်စားမယ့် ခေတ်ပြေပြေ ရေးပါ။ Emoji တွေ၊ လူငယ်စကားတွေ သုံးပါ။",
        "👔 Professional": "ကျွမ်းကျင်ပညာရှင်ဆန်ဆန်၊ တိကျစွာ ရေးပါ။"
    }

def get_durations():
    return {
        "⚡ 1 မိနစ် (~150 words)": 150,
        "📝 3 မိနစ် (~450 words)": 450,
        "📄 5 မိနစ် (~750 words)": 750,
        "📑 10 မိနစ် (~1500 words)": 1500,
        "📃 15 မိနစ် (~2250 words)": 2250,
        "📋 20 မိနစ် (~3000 words)": 3000,
        "📚 25 မိနစ် (~3750 words)": 3750,
        "📖 30 မိနစ် (~4500 words)": 4500,
        "📕 35 မိနစ် (~5250 words)": 5250,
        "📗 45 မိနစ် (~6750 words)": 6750,
        "📘 1 နာရီ (~9000 words)": 9000
    }

def get_content_prompt(ctype, title, words, tone="", custom=""):
    base = {
        "article": "အချက်အလက်ပြည့်စုံပြီး စိတ်ဝင်စားဖွယ်ကောင်းသော ဆောင်းပါး",
        "success": "လက်တွေ့ကျပြီး လှုံ့ဆော်နိုင်သော အောင်မြင်ရေးနည်းလမ်းများ",
        "story": "စိတ်ဝင်စားဖွယ်ကောင်းပြီး သင်ခန်းစာပါသော ဝတ္ထုတို",
        "tale": "ကလေးများနှင့် လူတိုင်းဖတ်ရှုနိုင်သော သင်ခန်းစာပါ ပုံပြင်",
        "news": "ဂျာနယ်လစ်ပုံစံ Who/What/When/Where/Why ပါဝင်သော သတင်း",
        "drama": "စိတ်လှုပ်ရှားဖွယ် dialog များပါဝင်သော ဇာတ်လမ်း",
        "horror": "တဖြည်းဖြည်း ထိတ်လန့်လာစေသော သရဲဇာတ်လမ်း",
        "tragic": "နက်နဲသော ခံစားချက်နှင့် ဘဝသင်ခန်းစာပါသော ဂမ္ဘီရဇာတ်လမ်း",
        "romance": "ရင်ခုန်ဖွယ်ကောင်းသော အချစ်ဇာတ်လမ်း",
        "fantasy": "မှော်ဆန်ပြီး စွန့်စားခန်းများပါသော စိတ်ကူးယဉ်ဇာတ်လမ်း",
        "mystery": "စုံထောက်ပုံစံ အံ့အားသင့်ဖွယ် အဆုံးသတ်ပါသော လျှို့ဝှက်ဆန်းကြယ်ဇာတ်လမ်း",
        "comedy": "ရယ်မောဖွယ်ကောင်းပြီး ပျော်ရွှင်စေသော ဟာသဇာတ်လမ်း",
        "motivational": "အားပေးစကားများနှင့် လုပ်ဆောင်နိုင်သော အကြံဉာဏ်များပါသော လှုံ့ဆော်စာ",
        "educational": "ရှင်းလင်းလွယ်ကူပြီး ဥပမာများပါသော ပညာရေးဆိုင်ရာ အကြောင်းအရာ",
        "custom": "အကြောင်းအရာ"
    }
    
    desc = base.get(ctype, base["article"])
    tone_inst = f"\n\nအရေးအသားပုံစံ: {tone}" if tone else ""
    custom_inst = f"\n\nအထူးညွှန်ကြားချက်: {custom}" if custom else ""
    
    return f"""သင်သည် မြန်မာစာ အရေးအသားကျွမ်းကျင်သူ ပရော်ဖက်ရှင်နယ် စာရေးဆရာတစ်ဦးဖြစ်သည်။

**တာဝန်**: "{title}" ခေါင်းစဉ်ဖြင့် {desc} ရေးပါ။

**လိုအပ်ချက်များ**:
1. စာလုံးရေ: {words} words ဝန်းကျင် (တိတိကျကျ ရေးပါ)
2. ဘာသာစကား: မြန်မာဘာသာ 100%
3. အရည်အသွေး: ပရော်ဖက်ရှင်နယ် content creator အဆင့်
4. ဖော်မတ်: TTS အတွက် သင့်တော်အောင် အပိုဒ်တိုများ၊ စာကြောင်းတိုများဖြင့် ရေးပါ
5. ဖတ်ရှုသူကို ဆွဲဆောင်နိုင်သော အဖွင့်စာပိုဒ်ဖြင့် စတင်ပါ
6. အဆုံးသတ်တွင် နိဂုံးချုပ် သို့မဟုတ် သင်ခန်းစာ ထည့်ပါ
{tone_inst}{custom_inst}

ယခုပဲ ရေးပါ။ မိတ်ဆက်စာ မလိုပါ။ Content တိုက်ရိုက် ရေးပါ။"""

# === EXPORT FUNCTIONS ===
def export_docx(content, title):
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document()
        doc.add_heading(title, 0)
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()
    except:
        return None

# === INITIALIZE SESSION STATE ===
def init_session():
    defaults = {
        'user_session': None,
        'api_type': 'app',
        'own_api_key': '',
        'content_result': None,
        'tts_audio': None,
        'generated_images': [],
        'editor_content': '',
        'content_history': []
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session()

# === GET ACTIVE API KEY ===
def get_active_api_key():
    if st.session_state.get('api_type') == 'own' and st.session_state.get('own_api_key'):
        return st.session_state['own_api_key']
    return get_app_api_key()

def is_own_api():
    return st.session_state.get('api_type') == 'own' and st.session_state.get('own_api_key')

# === LOGIN PAGE ===
if not st.session_state['user_session']:
    st.markdown("""
    <div style="text-align: center; padding: 1.5rem 0;">
        <h1 style="font-size: 2rem; margin-bottom: 0.3rem;">🎬 AI Studio Pro</h1>
        <p style="color: #94a3b8;">Content Creator's Toolkit</p>
    </div>
    """, unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["🔐 Login", "📝 Register"])
    
    with tab1:
        with st.container(border=True):
            with st.form("login_form"):
                email = st.text_input("📧 Email")
                password = st.text_input("🔑 Password", type="password")
                if st.form_submit_button("Login", use_container_width=True):
                    if email and password:
                        user, msg = login(email, password)
                        if user:
                            st.session_state['user_session'] = user
                            st.rerun()
                        else:
                            st.error(f"❌ {msg}")
                    else:
                        st.warning("Email နှင့် Password ထည့်ပါ")
    
    with tab2:
        with st.container(border=True):
            email2 = st.text_input("📧 Email", key="reg_email")
            pass1 = st.text_input("🔑 Password", type="password", key="reg_pass1")
            pass2 = st.text_input("🔑 Confirm Password", type="password", key="reg_pass2")
            
            if st.button("Register", use_container_width=True):
                if email2 and pass1 and pass2:
                    if pass1 != pass2:
                        st.error("Password များ မတူညီပါ")
                    else:
                        ok, msg = register(email2, pass1)
                        if ok:
                            st.success(f"✅ {msg}")
                        else:
                            st.error(f"❌ {msg}")
                else:
                    st.warning("အကွက်အားလုံး ဖြည့်ပါ")

else:
    # === MAIN APP ===
    user = st.session_state['user_session']
    
    # Header
    col1, col2 = st.columns([4, 1])
    with col1:
        st.markdown("<h1 style='font-size:1.5rem;margin:0;'>🎬 AI Studio Pro</h1>", unsafe_allow_html=True)
        st.caption(f"👤 {user['email']}")
    with col2:
        if st.button("🚪", help="Logout"):
            st.session_state['user_session'] = None
            st.rerun()
    
    # Admin Panel
    if user.get('is_admin'):
        with st.expander("🔧 Admin"):
            if supabase:
                users = supabase.table('users').select('*').order('created_at', desc=True).execute().data or []
                for u in users:
                    c1, c2, c3 = st.columns([3, 1, 1])
                    with c1:
                        st.write(u['email'])
                    with c2:
                        st.caption("✅" if u['approved'] else "⏳")
                    with c3:
                        if u['email'] != user['email']:
                            if st.button("Toggle", key=f"adm_{u['id']}"):
                                supabase.table('users').update({'approved': not u['approved']}).eq('id', u['id']).execute()
                                st.rerun()
    
    st.markdown("---")
    
    # === API SETTINGS ===
    with st.container(border=True):
        st.subheader("⚙️ API ရွေးချယ်မှု")
        
        api_type = st.radio(
            "API Type",
            ["🏢 App API (တရက် 10 ကြိမ်)", "🔑 Own API (Unlimited)"],
            horizontal=True,
            index=0 if st.session_state.get('api_type') == 'app' else 1
        )
        
        st.session_state['api_type'] = 'app' if 'App API' in api_type else 'own'
        
        if st.session_state['api_type'] == 'own':
            own_key = st.text_input(
                "🔑 သင့် Google AI API Key",
                type="password",
                value=st.session_state.get('own_api_key', ''),
                placeholder="AIza..."
            )
            st.session_state['own_api_key'] = own_key
            
            if own_key:
                try:
                    genai.configure(api_key=own_key)
                    st.success("✅ API ချိတ်ဆက်ပြီးပါပြီ - Unlimited အသုံးပြုနိုင်ပါတယ်")
                except:
                    st.error("❌ API Key မှားနေပါတယ်")
            else:
                st.info("💡 Own API သုံးရင် Unlimited ဖြစ်ပြီး Model အားလုံး ရွေးလို့ရပါတယ်")
        else:
            app_key = get_app_api_key()
            if app_key:
                try:
                    genai.configure(api_key=app_key)
                    # Show usage
                    usage = get_usage(user['id'])
                    st.markdown(f"""
                    <div class="usage-box">
                        <b>📊 ယနေ့ အသုံးပြုမှု:</b><br>
                        ✍️ Content: {usage['content']}/10 | 
                        🌐 Translate: {usage['translate']}/10 | 
                        🎙️ TTS: {usage['tts']}/10 | 
                        🖼️ Thumbnail: {usage['thumbnail']}/10
                    </div>
                    """, unsafe_allow_html=True)
                    st.caption("💡 App API သုံးရင် Feature တခုချင်းစီ တရက် 10 ကြိမ်စီ သုံးနိုင်ပါတယ်။ Edge TTS က Unlimited ပါ။")
                except:
                    st.error("❌ App API ချိတ်ဆက်မှု မအောင်မြင်ပါ")
            else:
                st.error("❌ App API Key မရှိပါ။ Own API သုံးပါ။")
    
    st.markdown("---")
    
    # === MAIN TABS ===
    tab1, tab2, tab3, tab4 = st.tabs(["✍️ Content", "🌐 Translate", "🎙️ TTS", "🖼️ Thumbnail"])
    
    # === TAB 1: CONTENT ===
    with tab1:
        st.header("✍️ Content Creator")
        
        # Check limit
        can_use, remaining = check_limit(user['id'], 'content', st.session_state['api_type'])
        
        if not can_use and st.session_state['api_type'] == 'app':
            st.markdown("""
            <div class="limit-warning">
                ⚠️ ယနေ့အတွက် Content limit (10 ကြိမ်) ပြည့်သွားပါပြီ။<br>
                မနက်ဖြန် ပြန်သုံးနိုင်ပါတယ် သို့မဟုတ် Own API သုံးပါ။
            </div>
            """, unsafe_allow_html=True)
        
        with st.container(border=True):
            title = st.text_input("📝 ခေါင်းစဉ်", placeholder="ဥပမာ: ဘဝမှာ အောင်မြင်ဖို့ လိုအပ်တဲ့ အရာ ၅ ခု")
            
            col1, col2 = st.columns(2)
            with col1:
                ctype = st.selectbox("📂 အမျိုးအစား", list(get_content_types().keys()))
            with col2:
                duration = st.selectbox("⏱️ ကြာချိန်", list(get_durations().keys()))
            
            tone = st.selectbox("🎨 အရေးအသားပုံစံ", list(get_tones().keys()))
            
            # Model selection based on API type
            # App API = All models
            # Own API = User ရွေးလို့ရ + Free Tier warning

            all_models = [
                "models/gemini-2.5-flash",
                "models/gemini-2.5-pro",
                "gemini-2.0-flash-exp",
                "gemini-1.5-flash"
            ]

            model = st.selectbox("🤖 Model", all_models)

            # Own API + Pro model selected = show warning
            if is_own_api() and "pro" in model.lower():
                st.warning("⚠️ Pro model ကို Free Tier API နဲ့ သုံးရင် Rate Limit Error ရနိုင်ပါတယ်။ Billing Enabled API key သုံးပါ။")

            with st.expander("🎯 စိတ်ကြိုက် ညွှန်ကြားချက်"):
                custom = st.text_area("", placeholder="ဥပမာ: Emoji တွေထည့်ပေး...", height=80)
            
            if st.button("✨ Generate", use_container_width=True, type="primary", disabled=(not can_use and st.session_state['api_type'] == 'app')):
                api_key = get_active_api_key()
                if not api_key:
                    st.error("❌ API Key မရှိပါ")
                elif not title.strip():
                    st.warning("⚠️ ခေါင်းစဉ် ထည့်ပါ")
                else:
                    with st.spinner("✍️ Content ရေးနေပါတယ်..."):
                        try:
                            genai.configure(api_key=api_key)
                            m = genai.GenerativeModel(model)
                            
                            ctype_val = get_content_types()[ctype]
                            words = get_durations()[duration]
                            tone_val = get_tones()[tone]
                            
                            prompt = get_content_prompt(ctype_val, title, words, tone_val, custom)
                            
                            resp, err = call_api(m, prompt, 600)
                            
                            if resp:
                                result, _ = get_text(resp)
                                if result:
                                    st.session_state['content_result'] = result
                                    st.session_state['editor_content'] = result
                                    
                                    # Increment usage for APP API
                                    if st.session_state['api_type'] == 'app':
                                        increment_usage(user['id'], 'content')
                                    
                                    st.success("✅ ပြီးပါပြီ!")
                                    st.rerun()
                            else:
                                st.error(f"❌ {err}")
                        except Exception as e:
                            st.error(f"❌ {str(e)}")
        
        # Show result
        if st.session_state.get('content_result'):
            with st.container(border=True):
                st.subheader("📄 ရလဒ်")
                result = st.session_state['content_result']
                
                wc = len(result.split())
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("စာလုံးရေ", f"{wc:,}")
                with col2:
                    st.metric("ဖတ်ချိန်", f"~{max(1,wc//200)} min")
                with col3:
                    st.metric("ပြောချိန်", f"~{max(1,wc//150)} min")
                
                st.text_area("Content", result, height=300, label_visibility="collapsed")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.download_button("📄 TXT", result, f"{title[:20]}.txt", use_container_width=True)
                with col2:
                    if DOCX_AVAILABLE:
                        docx = export_docx(result, title)
                        if docx:
                            st.download_button("📝 DOCX", docx, f"{title[:20]}.docx", use_container_width=True)
                with col3:
                    if st.button("🗑️ Clear", use_container_width=True):
                        st.session_state['content_result'] = None
                        st.rerun()
    
    # === TAB 2: TRANSLATE ===
    with tab2:
        st.header("🌐 Translator")
        
        can_use, remaining = check_limit(user['id'], 'translate', st.session_state['api_type'])
        
        if not can_use and st.session_state['api_type'] == 'app':
            st.markdown("""
            <div class="limit-warning">
                ⚠️ ယနေ့အတွက် Translate limit (10 ကြိမ်) ပြည့်သွားပါပြီ။
            </div>
            """, unsafe_allow_html=True)
        
        st.info("💡 Google Drive link သို့မဟုတ် File upload က အဆင်အပြေဆုံးဖြစ်ပါတယ်။ YouTube public videos လည်း သုံးလို့ရပါတယ်။")
        
        with st.container(border=True):
            languages = {
                "🇲🇲 မြန်မာ": "Burmese",
                "🇺🇸 English": "English", 
                "🇹🇭 ไทย": "Thai",
                "🇨🇳 中文": "Chinese",
                "🇯🇵 日本語": "Japanese",
                "🇰🇷 한국어": "Korean"
            }
            
            col1, col2 = st.columns([2, 1])
            with col1:
                target = st.selectbox("🎯 ဘာသာပြန်မည့်ဘာသာ", list(languages.keys()))
            with col2:
                with col2:
                if is_own_api():
                    trans_model = st.selectbox("Model", ["models/gemini-2.5-flash", "models/gemini-2.5-pro"], key="tm")
                    if "pro" in trans_model.lower():
                        st.warning("⚠️ Pro model ကို Free Tier API နဲ့ သုံးရင် Rate Limit Error ရနိုင်ပါတယ်။")
                else:
                    trans_model = "models/gemini-2.5-flash"
                    st.caption("Flash model")
            
            input_type = st.radio(
                "Input နည်းလမ်း",
                ["📋 Text Paste", "📤 File Upload", "🔗 URL"],
                horizontal=True
            )
            
            text_input = None
            file_input = None
            url_input = None
            
            if input_type == "📋 Text Paste":
                text_input = st.text_area("ဘာသာပြန်မည့် စာသား", height=150, placeholder="ဒီမှာ စာသား paste လုပ်ပါ...")
            elif input_type == "📤 File Upload":
                file_input = st.file_uploader("File", type=["mp3", "mp4", "txt", "srt", "docx"])
            else:
                url_input = st.text_input("🔗 Video URL", placeholder="YouTube, TikTok, Google Drive link...")
            
            if st.button("🌐 Translate", use_container_width=True, type="primary", disabled=(not can_use and st.session_state['api_type'] == 'app')):
                api_key = get_active_api_key()
                if not api_key:
                    st.error("❌ API Key မရှိပါ")
                elif not text_input and not file_input and not url_input:
                    st.warning("⚠️ Input ထည့်ပါ")
                else:
                    target_lang = languages[target]
                    genai.configure(api_key=api_key)
                    model = genai.GenerativeModel(trans_model)
                    
                    # Text translation
                    if text_input:
                        with st.spinner("🌐 ဘာသာပြန်နေပါတယ်..."):
                            resp, err = call_api(model, f"Translate to {target_lang}. Return ONLY translated text:\n\n{text_input}", 300)
                            if resp:
                                result, _ = get_text(resp)
                                if result:
                                    if st.session_state['api_type'] == 'app':
                                        increment_usage(user['id'], 'translate')
                                    st.text_area("ရလဒ်", result, height=200)
                                    col1, col2 = st.columns(2)
                                    with col1:
                                        st.download_button("📄 TXT", result, "translated.txt", use_container_width=True)
                                    with col2:
                                        st.download_button("📋 SRT", text_to_srt(result), "translated.srt", use_container_width=True)
                            else:
                                st.error(f"❌ {err}")
                    
                    # File translation
                    elif file_input:
                        ext = file_input.name.split('.')[-1].lower()
                        
                        if ext in ['txt', 'srt']:
                            txt = file_input.getvalue().decode('utf-8')
                            with st.spinner("🌐 ဘာသာပြန်နေပါတယ်..."):
                                resp, err = call_api(model, f"Translate to {target_lang}:\n\n{txt}", 600)
                                if resp:
                                    result, _ = get_text(resp)
                                    if result:
                                        if st.session_state['api_type'] == 'app':
                                            increment_usage(user['id'], 'translate')
                                        st.text_area("ရလဒ်", result, height=200)
                                        st.download_button("📄 Download", result, f"trans_{file_input.name}", use_container_width=True)
                                else:
                                    st.error(f"❌ {err}")
                        
                        elif ext in ['mp3', 'mp4']:
                            progress = st.progress(0)
                            status = st.empty()
                            
                            status.info("📤 Uploading...")
                            progress.progress(20)
                            
                            path, _ = save_up(file_input)
                            if path:
                                progress.progress(40)
                                gem_file = upload_gem(path)
                                
                                if gem_file:
                                    status.info("🌐 ဘာသာပြန်နေပါတယ်...")
                                    progress.progress(60)
                                    
                                    resp, err = call_api(model, [gem_file, f"Transcribe and translate to {target_lang}. Return ONLY translated text."], 900)
                                    progress.progress(90)
                                    
                                    if resp:
                                        result, _ = get_text(resp)
                                        progress.progress(100)
                                        status.success("✅ ပြီးပါပြီ!")
                                        
                                        if result:
                                            if st.session_state['api_type'] == 'app':
                                                increment_usage(user['id'], 'translate')
                                            st.text_area("ရလဒ်", result, height=200)
                                            col1, col2 = st.columns(2)
                                            with col1:
                                                st.download_button("📄 TXT", result, "translated.txt", use_container_width=True)
                                            with col2:
                                                st.download_button("📋 SRT", text_to_srt(result), "translated.srt", use_container_width=True)
                                    else:
                                        status.error(f"❌ {err}")
                                    
                                    try:
                                        genai.delete_file(gem_file.name)
                                    except:
                                        pass
                                else:
                                    status.error("❌ Upload မအောင်မြင်ပါ")
                                rm_file(path)
                    
                    # URL translation
                    elif url_input:
                        progress = st.progress(0)
                        status = st.empty()
                        
                        progress.progress(10)
                        path, err = download_video_url(url_input, status)
                        
                        if path:
                            progress.progress(30)
                            status.info("📤 Uploading...")
                            
                            gem_file = upload_gem(path)
                            
                            if gem_file:
                                status.info("🌐 ဘာသာပြန်နေပါတယ်...")
                                progress.progress(60)
                                
                                resp, err = call_api(model, [gem_file, f"Transcribe and translate to {target_lang}. Return ONLY translated text."], 900)
                                progress.progress(90)
                                
                                if resp:
                                    result, _ = get_text(resp)
                                    progress.progress(100)
                                    status.success("✅ ပြီးပါပြီ!")
                                    
                                    if result:
                                        if st.session_state['api_type'] == 'app':
                                            increment_usage(user['id'], 'translate')
                                        st.text_area("ရလဒ်", result, height=200)
                                        col1, col2 = st.columns(2)
                                        with col1:
                                            st.download_button("📄 TXT", result, "translated.txt", use_container_width=True)
                                        with col2:
                                            st.download_button("📋 SRT", text_to_srt(result), "translated.srt", use_container_width=True)
                                else:
                                    status.error(f"❌ {err}")
                                
                                try:
                                    genai.delete_file(gem_file.name)
                                except:
                                    pass
                            else:
                                status.error("❌ Upload မအောင်မြင်ပါ")
                            rm_file(path)
                        else:
                            status.error(f"❌ {err}")
    
    # === TAB 3: TTS ===
    with tab3:
        st.header("🎙️ Text to Speech")
        
        with st.container(border=True):
            engine = st.radio("Engine", ["⚡ Edge TTS (Unlimited)", "✨ Gemini TTS"], horizontal=True)
            
            st.markdown("---")
            
            if "Edge" in engine:
                if not EDGE_TTS_AVAILABLE:
                    st.error("❌ Edge TTS မရနိုင်ပါ")
                else:
                    tts_text = st.text_area("📝 စာသား", height=180, placeholder="ဒီမှာ စာသား ရိုက်ထည့်ပါ...")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        voice = st.selectbox("🔊 Voice", list(edge_voices().keys()))
                    with col2:
                        rate = st.slider("⚡ Speed", -50, 50, 0, format="%d%%")
                    
                    st.caption(f"📊 {len(tts_text)} characters")
                    
                    if st.button("🎙️ Generate Audio", use_container_width=True, type="primary"):
                        if tts_text.strip():
                            with st.spinner("🔄 Audio ထုတ်နေပါတယ်..."):
                                path, err = gen_edge(tts_text, edge_voices()[voice], rate)
                                if path:
                                    st.session_state['tts_audio'] = path
                                    st.success("✅ ပြီးပါပြီ!")
                                    st.rerun()
                                else:
                                    st.error(f"❌ {err}")
                        else:
                            st.warning("⚠️ စာသား ထည့်ပါ")
            
            else:  # Gemini TTS
                can_use, _ = check_limit(user['id'], 'tts', st.session_state['api_type'])
                
                if not can_use and st.session_state['api_type'] == 'app':
                    st.markdown("""
                    <div class="limit-warning">
                        ⚠️ ယနေ့အတွက် Gemini TTS limit (10 ကြိမ်) ပြည့်သွားပါပြီ။ Edge TTS က Unlimited သုံးလို့ရပါတယ်။
                    </div>
                    """, unsafe_allow_html=True)
                
                if not GENAI_NEW_AVAILABLE:
                    st.error("❌ Gemini TTS မရနိုင်ပါ")
                else:
                    tts_text = st.text_area("📝 စာသား", height=180, placeholder="ဒီမှာ စာသား ရိုက်ထည့်ပါ...", key="gem_txt")
                    
                    style = st.selectbox("🎨 Voice Style", list(voice_styles().keys()))
                    style_prompt = voice_styles()[style]
                    
                    if "စိတ်ကြိုက်" in style:
                        style_prompt = st.text_input("Custom style", placeholder="Describe voice style...")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        voice = st.selectbox("🔊 Voice", list(gemini_voices().keys()))
                    with col2:
                        tts_model = st.selectbox("Model", ["gemini-2.5-flash-preview-tts", "gemini-2.5-pro-preview-tts"])
                    with col3:
                        speed = st.slider("Speed", 0.5, 2.0, 1.0, 0.1)
                    
                    st.caption(f"📊 {len(tts_text)} characters")
                    
                    if st.button("🎙️ Generate Audio", use_container_width=True, type="primary", key="gem_btn", disabled=(not can_use and st.session_state['api_type'] == 'app')):
                        api_key = get_active_api_key()
                        if not api_key:
                            st.error("❌ API Key မရှိပါ")
                        elif not tts_text.strip():
                            st.warning("⚠️ စာသား ထည့်ပါ")
                        else:
                            with st.spinner("🔄 Audio ထုတ်နေပါတယ်..."):
                                path, err = gen_gemini_tts(api_key, tts_text, gemini_voices()[voice], tts_model, style_prompt, speed)
                                if path:
                                    if st.session_state['api_type'] == 'app':
                                        increment_usage(user['id'], 'tts')
                                    st.session_state['tts_audio'] = path
                                    st.success("✅ ပြီးပါပြီ!")
                                    st.rerun()
                                else:
                                    st.error(f"❌ {err}")
        
        # Audio output
        if st.session_state.get('tts_audio') and os.path.exists(st.session_state['tts_audio']):
            with st.container(border=True):
                st.subheader("🔊 Audio")
                with open(st.session_state['tts_audio'], 'rb') as f:
                    audio_bytes = f.read()
                
                mime = "audio/wav" if st.session_state['tts_audio'].endswith(".wav") else "audio/mp3"
                st.audio(audio_bytes, format=mime)
                
                ext = "wav" if ".wav" in st.session_state['tts_audio'] else "mp3"
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button("📥 Download", audio_bytes, f"audio.{ext}", use_container_width=True)
                with col2:
                    if st.button("🗑️ Clear", use_container_width=True):
                        rm_file(st.session_state['tts_audio'])
                        st.session_state['tts_audio'] = None
                        st.rerun()
    
    # === TAB 4: THUMBNAIL ===
    with tab4:
        st.header("🖼️ AI Thumbnail")
        
        can_use, _ = check_limit(user['id'], 'thumbnail', st.session_state['api_type'])
        
        if not can_use and st.session_state['api_type'] == 'app':
            st.markdown("""
            <div class="limit-warning">
                ⚠️ ယနေ့အတွက် Thumbnail limit (10 ကြိမ်) ပြည့်သွားပါပြီ။
            </div>
            """, unsafe_allow_html=True)
        
        with st.container(border=True):
            # Reference images
            ref_imgs = st.file_uploader("🖼️ Reference Images (Optional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
            
            if ref_imgs:
                cols = st.columns(min(len(ref_imgs), 5))
                for i, img in enumerate(ref_imgs[:5]):
                    with cols[i]:
                        st.image(img, use_container_width=True)
            
            st.markdown("---")
            
            # Templates
            templates = {
                "🎨 Custom": "",
                "🎬 Movie Recap": "dramatic YouTube movie recap thumbnail, cinematic lighting, emotional scene, bold title text",
                "😱 Shocking": "YouTube thumbnail, shocked expression, bright red yellow background, dramatic text",
                "👻 Horror": "horror movie thumbnail, dark scary atmosphere, creepy shadows",
                "💕 Romance": "romantic thumbnail, soft pink lighting, dreamy bokeh"
            }
            
            template = st.selectbox("📋 Template", list(templates.keys()))
            
            # Model selection
            # App API = All models
# Own API = User ရွေးလို့ရ + Free Tier warning

all_models = [
    "models/gemini-2.5-flash",
    "models/gemini-2.5-pro", 
    "gemini-2.0-flash-exp",
    "gemini-1.5-flash"
]

model = st.selectbox("🤖 Model", all_models)

# Own API + Pro model selected = show warning
if is_own_api() and "pro" in model.lower():
    st.warning("⚠️ Pro model ကို Free Tier API နဲ့ သုံးရင် Rate Limit Error ရနိုင်ပါတယ်။ Billing Enabled API key သုံးပါ။")
            
            # Size
            sizes = {
                "📺 16:9 (1280x720)": "1280x720",
                "📱 9:16 (720x1280)": "720x1280",
                "⬜ 1:1 (1024x1024)": "1024x1024",
                "🖼️ 4:3 (1024x768)": "1024x768",
                "📷 3:4 (768x1024)": "768x1024"
            }
            size = st.selectbox("📐 Size", list(sizes.keys()))
            
            # Prompt
            prompt = st.text_area("✏️ Prompt", value=templates[template], height=80, placeholder="Describe your thumbnail...")
            
            # Text and negative prompt
            col1, col2 = st.columns(2)
            with col1:
                add_text = st.text_input("📝 Add Text", placeholder="EP.1, Part 2...")
            with col2:
                neg_prompt = st.text_input("🚫 Negative Prompt", placeholder="blurry, low quality...")
            
            # Count
            num_imgs = st.selectbox("🔢 Count", [1, 2, 3, 4])
            
            if st.button("✨ Generate", use_container_width=True, type="primary", disabled=(not can_use and st.session_state['api_type'] == 'app')):
                api_key = get_active_api_key()
                if not api_key:
                    st.error("❌ API Key မရှိပါ")
                elif not prompt.strip():
                    st.warning("⚠️ Prompt ထည့်ပါ")
                else:
                    st.session_state['generated_images'] = []
                    
                    # Build prompt
                    size_val = sizes[size]
                    final_prompt = prompt.strip()
                    if add_text:
                        final_prompt += f", text saying '{add_text}', bold text"
                    final_prompt += f", {size_val}, high quality"
                    if neg_prompt:
                        final_prompt += f". Avoid: {neg_prompt}"
                    
                    # Load reference images
                    ref_pil = []
                    if ref_imgs:
                        for r in ref_imgs[:5]:
                            try:
                                r.seek(0)
                                ref_pil.append(Image.open(io.BytesIO(r.read())))
                            except:
                                pass
                    
                    genai.configure(api_key=api_key)
                    
                    progress = st.progress(0)
                    status = st.empty()
                    
                    for i in range(1, num_imgs + 1):
                        status.info(f"🎨 Generating {i}/{num_imgs}...")
                        
                        try:
                            mdl = genai.GenerativeModel(thumb_model)
                            content = [f"Generate image: {final_prompt}"]
                            if ref_pil:
                                content.extend(ref_pil)
                            
                            resp = mdl.generate_content(content, request_options={"timeout": 300})
                            
                            if resp.candidates:
                                for p in resp.candidates[0].content.parts:
                                    if hasattr(p, 'inline_data') and p.inline_data:
                                        img_data = p.inline_data.data
                                        if img_data and len(img_data) > 1000:
                                            st.session_state['generated_images'].append({
                                                'data': img_data,
                                                'mime': p.inline_data.mime_type,
                                                'idx': i
                                            })
                                            status.success(f"✅ Image {i} done!")
                                            break
                        except Exception as e:
                            status.warning(f"⚠️ Image {i}: {str(e)[:50]}")
                        
                        progress.progress(i / num_imgs)
                        if i < num_imgs:
                            time.sleep(1)
                    
                    if st.session_state['generated_images']:
                        if st.session_state['api_type'] == 'app':
                            increment_usage(user['id'], 'thumbnail')
                        status.success(f"✅ {len(st.session_state['generated_images'])}/{num_imgs} images generated!")
                    else:
                        status.error("❌ မအောင်မြင်ပါ")
            
            # Show results
            if st.session_state.get('generated_images'):
                st.markdown("---")
                st.subheader("🖼️ Results")
                
                if st.button("🗑️ Clear All"):
                    st.session_state['generated_images'] = []
                    st.rerun()
                
                for i, img in enumerate(st.session_state['generated_images']):
                    with st.container(border=True):
                        st.image(img['data'], use_container_width=True)
                        st.download_button(
                            f"📥 Download #{img['idx']}",
                            img['data'],
                            f"thumbnail_{img['idx']}.png",
                            mime=img.get('mime', 'image/png'),
                            key=f"thumb_dl_{img['idx']}",
                            use_container_width=True
                        )
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align:center;padding:0.5rem;">
        <p style="color:#64748b;font-size:0.8rem;">🎬 AI Studio Pro v8.0 | Made for Content Creators</p>
    </div>
    """, unsafe_allow_html=True)
